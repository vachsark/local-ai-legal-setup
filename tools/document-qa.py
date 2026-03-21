"""
title: Document Q&A
author: local-ai-legal-setup
version: 0.2.0
license: MIT
description: Upload a legal document (PDF, DOCX, TXT) and ask questions about it. Returns cited answers pinpointing the exact section and paragraph where the answer was found. Never fabricates — says "not found in document" when the answer is absent. Uses hybrid BM25 + semantic retrieval (Ollama embeddings) when available — falls back to BM25 if no embedding model is installed.
"""

import json
import math
import re
import urllib.request
import urllib.error
from typing import Optional

from pydantic import BaseModel


# ── Embedding helpers (no external deps — pure stdlib) ──────────────────────

def _get_embedding(text: str, model: str) -> Optional[list]:
    """
    Request an embedding vector from Ollama's local API.
    Returns a list of floats, or None if the request fails (model not available,
    Ollama not running, etc.). Caller should fall back to BM25 on None.
    """
    try:
        payload = json.dumps({"model": model, "prompt": text}).encode()
        req = urllib.request.Request(
            "http://localhost:11434/api/embeddings",
            data=payload,
            headers={"Content-Type": "application/json"},
            method="POST",
        )
        with urllib.request.urlopen(req, timeout=30) as resp:
            data = json.loads(resp.read())
            return data.get("embedding")
    except Exception:
        return None


def _cosine_similarity(a: list, b: list) -> float:
    """Compute cosine similarity between two equal-length vectors."""
    dot = sum(x * y for x, y in zip(a, b))
    norm_a = sum(x * x for x in a) ** 0.5
    norm_b = sum(x * x for x in b) ** 0.5
    if norm_a == 0 or norm_b == 0:
        return 0.0
    return dot / (norm_a * norm_b)


def _detect_embedding_model() -> Optional[str]:
    """
    Ask Ollama which embedding models are available.
    Preference order: nomic-embed-text → qwen3-embedding → any *embed* model.
    Returns None if Ollama isn't running or no embedding model is found.
    """
    try:
        req = urllib.request.Request(
            "http://localhost:11434/api/tags",
            method="GET",
        )
        with urllib.request.urlopen(req, timeout=5) as resp:
            data = json.loads(resp.read())
            models = [m["name"] for m in data.get("models", [])]

        # Ranked preference
        for preferred in ("nomic-embed-text", "qwen3-embedding:0.6b", "qwen3-embedding"):
            for m in models:
                if m.startswith(preferred):
                    return m
        # Any embed-named model as fallback
        for m in models:
            if "embed" in m.lower():
                return m
        return None
    except Exception:
        return None


# ── Reciprocal Rank Fusion ──────────────────────────────────────────────────

def _rrf_fuse(bm25_ranked: list, sem_ranked: list, k: int = 60) -> list:
    """
    Combine two ranked lists using Reciprocal Rank Fusion.
    Each item gets score = 1/(k + rank) from each list; scores are summed.
    Returns items sorted by fused score descending, with 'score' field updated.
    """
    scores: dict[int, float] = {}

    for rank, chunk in enumerate(bm25_ranked, start=1):
        idx = chunk["_idx"]
        scores[idx] = scores.get(idx, 0.0) + 1.0 / (k + rank)

    for rank, chunk in enumerate(sem_ranked, start=1):
        idx = chunk["_idx"]
        scores[idx] = scores.get(idx, 0.0) + 1.0 / (k + rank)

    # Collect all chunks that appeared in either list
    seen: dict[int, dict] = {}
    for chunk in bm25_ranked + sem_ranked:
        idx = chunk["_idx"]
        if idx not in seen:
            seen[idx] = chunk

    fused = [
        {**chunk, "score": scores[chunk["_idx"]]}
        for chunk in seen.values()
    ]
    fused.sort(key=lambda x: x["score"], reverse=True)
    return fused


class Tools:
    class Valves(BaseModel):
        chunk_size: int = 1500
        overlap: int = 200
        max_chunks: int = 10  # base; auto-scales for large docs (see _retrieve)
        max_chunks_cap: int = 20  # hard upper bound regardless of doc size
        max_doc_chars: int = 200000
        embedding_model: str = ""  # leave blank for auto-detect
        semantic_weight: float = 0.5  # RRF balance (unused — RRF is rank-based)

    def __init__(self):
        self.valves = self.Valves()
        # Per-instance embedding cache: chunk text → vector
        # Populated once per document; cleared when a new doc is processed.
        self._embed_cache: dict[str, list] = {}
        self._embed_model: Optional[str] = None  # resolved at first embed call
        self._embed_available: Optional[bool] = None  # None = not yet probed

    async def query_document(
        self,
        query: str,
        __files__: Optional[list] = None,
        __event_emitter__=None,
    ) -> str:
        """
        Ask a question about an uploaded legal document. Returns an answer
        with citations to the specific section and paragraph where the
        information was found.

        Upload a PDF, DOCX, or TXT file and then ask your question.
        Uses hybrid BM25 + semantic search when an Ollama embedding model is
        available — falls back to BM25-only if not.

        :param query: The question to answer based on the document.
        :return: Relevant document excerpts with section citations for the LLM to answer from.
        """
        if __event_emitter__:
            await __event_emitter__(
                {
                    "type": "status",
                    "data": {
                        "description": "Loading document...",
                        "done": False,
                    },
                }
            )

        if not __files__ or len(__files__) == 0:
            return (
                "No document uploaded. Please attach a PDF, DOCX, or TXT file "
                "to your message using the paperclip icon, then ask your question."
            )

        if not query or not query.strip():
            return "Please provide a question to answer from the document."

        file_obj = __files__[0]
        filename = file_obj.get("filename", file_obj.get("name", "document"))

        try:
            text = self._extract_text(file_obj)
        except Exception as e:
            if __event_emitter__:
                await __event_emitter__(
                    {
                        "type": "status",
                        "data": {"description": "Failed to read file", "done": True},
                    }
                )
            return f"Error reading file '{filename}': {str(e)}"

        if not text or not text.strip():
            return f"The document '{filename}' appears to be empty or could not be read."

        # Truncate if enormous
        if len(text) > self.valves.max_doc_chars:
            text = text[: self.valves.max_doc_chars]

        if __event_emitter__:
            await __event_emitter__(
                {
                    "type": "status",
                    "data": {
                        "description": f"Chunking document ({len(text):,} chars)...",
                        "done": False,
                    },
                }
            )

        # Chunk and index
        chunks = self._chunk_document(text)

        # Probe for embedding model once per instance
        if self._embed_available is None:
            model_override = self.valves.embedding_model.strip()
            if model_override:
                self._embed_model = model_override
                probe = _get_embedding("test", self._embed_model)
                self._embed_available = probe is not None
            else:
                self._embed_model = _detect_embedding_model()
                self._embed_available = self._embed_model is not None

        retrieval_mode = "hybrid" if self._embed_available else "BM25"

        if __event_emitter__:
            await __event_emitter__(
                {
                    "type": "status",
                    "data": {
                        "description": (
                            f"Searching {len(chunks)} sections "
                            f"({retrieval_mode}) for: {query[:50]}..."
                        ),
                        "done": False,
                    },
                }
            )

        # Scale max_chunks for large documents: proportional to chunk count,
        # floored at the base valve and capped at max_chunks_cap.
        scaled_max = max(
            self.valves.max_chunks,
            min(self.valves.max_chunks_cap, len(chunks) // 15),
        )

        # Retrieve relevant chunks
        top_chunks = self._retrieve(query, chunks, scaled_max)

        if not top_chunks:
            if __event_emitter__:
                await __event_emitter__(
                    {
                        "type": "status",
                        "data": {
                            "description": "No relevant sections found",
                            "done": True,
                        },
                    }
                )
            return json.dumps(
                {
                    "document": filename,
                    "query": query,
                    "result": "not_found",
                    "answer_instruction": (
                        "No relevant sections were found for this query. "
                        "Answer: 'This information was not found in the document.'"
                    ),
                    "chunks": [],
                },
                indent=2,
            )

        if __event_emitter__:
            await __event_emitter__(
                {
                    "type": "status",
                    "data": {
                        "description": f"Found {len(top_chunks)} relevant section(s) [{retrieval_mode}]",
                        "done": True,
                    },
                }
            )

        output = {
            "document": filename,
            "query": query,
            "retrieval_mode": retrieval_mode,
            "embedding_model": self._embed_model if self._embed_available else None,
            "total_sections_searched": len(chunks),
            "relevant_sections_found": len(top_chunks),
            "answer_instruction": (
                "Answer the query ONLY using the document excerpts below. "
                "For every fact in your answer, cite the section using the "
                "'citation' field (e.g., 'See Section 3.2, paragraph 2'). "
                "If the answer is not in these excerpts, say: "
                "'This information was not found in the document.' "
                "Never fabricate or infer beyond what the text states."
            ),
            "excerpts": [
                {
                    "citation": chunk["citation"],
                    "score": round(chunk["score"], 4),
                    "text": chunk["text"],
                }
                for chunk in top_chunks
            ],
        }

        return json.dumps(output, indent=2)

    # ── Text extraction ────────────────────────────────────────────────────────

    def _extract_text(self, file_obj: dict) -> str:
        """Extract plain text from a file object (supports TXT, basic DOCX, PDF via base64)."""
        filename = file_obj.get("filename", file_obj.get("name", ""))

        # Open WebUI passes file content in various shapes depending on version
        content = None
        if "content" in file_obj:
            content = file_obj["content"]
        elif "data" in file_obj and isinstance(file_obj["data"], dict):
            content = file_obj["data"].get("content", "")
        elif "data" in file_obj and isinstance(file_obj["data"], str):
            content = file_obj["data"]
        elif "path" in file_obj:
            return self._read_file_from_path(file_obj["path"], filename)

        if content is None:
            raise ValueError(
                f"Could not extract text from '{filename}'. "
                "Ensure the file is TXT or that Open WebUI has extracted its content."
            )

        # If it looks like a base64 blob, try to decode
        if isinstance(content, str) and self._is_base64(content):
            content = self._decode_base64_text(content, filename)

        return content if isinstance(content, str) else str(content)

    def _read_file_from_path(self, path: str, filename: str) -> str:
        """Read text from a file path, handling PDF via pdftotext and DOCX via python-docx."""
        import subprocess

        ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

        if ext == "pdf":
            try:
                result = subprocess.run(
                    ["pdftotext", path, "-"],
                    capture_output=True,
                    text=True,
                    timeout=30,
                )
                if result.returncode == 0:
                    return result.stdout
                raise ValueError(
                    f"pdftotext failed (exit {result.returncode}): {result.stderr[:200]}"
                )
            except FileNotFoundError:
                raise ValueError(
                    "pdftotext not found. Install poppler-utils: "
                    "'brew install poppler' (Mac) or 'apt install poppler-utils' (Linux)."
                )

        elif ext == "docx":
            try:
                from docx import Document  # type: ignore
                doc = Document(path)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                paragraphs.append(cell.text)
                return "\n".join(paragraphs)
            except ImportError:
                raise ValueError(
                    "python-docx not installed. Run: pip install python-docx"
                )
            except Exception as e:
                raise ValueError(f"DOCX read error: {e}")

        else:
            with open(path, "r", errors="replace") as f:
                return f.read()

    def _is_base64(self, s: str) -> bool:
        """Quick heuristic — base64 strings are long, no whitespace, all valid chars."""
        if len(s) < 100:
            return False
        sample = s[:200].strip()
        b64_chars = set(
            "ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789+/="
        )
        return all(c in b64_chars for c in sample if c not in "\n\r")

    def _decode_base64_text(self, encoded: str, filename: str) -> str:
        """Decode base64 content to string. Handles PDF, DOCX, and plain text."""
        import base64
        import subprocess
        import tempfile
        import os

        data = base64.b64decode(encoded)
        ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

        if ext == "pdf":
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                tmp.write(data)
                tmp_path = tmp.name
            try:
                result = subprocess.run(
                    ["pdftotext", tmp_path, "-"],
                    capture_output=True,
                    text=True,
                    timeout=30,
                )
                if result.returncode == 0:
                    return result.stdout
                raise ValueError(
                    f"pdftotext failed: {result.stderr[:200]}"
                )
            except FileNotFoundError:
                raise ValueError(
                    "pdftotext not found. Install poppler-utils: "
                    "'brew install poppler' (Mac) or 'apt install poppler-utils' (Linux)."
                )
            finally:
                os.unlink(tmp_path)

        elif ext == "docx":
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp.write(data)
                tmp_path = tmp.name
            try:
                from docx import Document  # type: ignore
                doc = Document(tmp_path)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                paragraphs.append(cell.text)
                return "\n".join(paragraphs)
            except ImportError:
                raise ValueError(
                    "python-docx not installed. Run: pip install python-docx"
                )
            except Exception as e:
                raise ValueError(f"DOCX read error: {e}")
            finally:
                os.unlink(tmp_path)

        else:
            return data.decode("utf-8", errors="replace")

    # ── Chunking ───────────────────────────────────────────────────────────────

    def _chunk_document(self, text: str) -> list:
        """
        Split document into overlapping chunks, tracking section headers for citation.
        Returns list of dicts: {text, citation, char_start, char_end, _idx}.
        _idx is a stable integer index used for RRF deduplication.
        """
        lines = text.splitlines()
        chunks = []
        current_section = "Document"
        paragraph_num = 0

        def flush_chunk(lines_so_far: list, section: str, para_num: int):
            chunk_text = "\n".join(lines_so_far).strip()
            if len(chunk_text) > 50:  # skip micro-chunks
                chunks.append(
                    {
                        "text": chunk_text,
                        "citation": f"{section}, paragraph {para_num}",
                        "section": section,
                        "paragraph": para_num,
                        "_idx": len(chunks),
                    }
                )

        buffer_lines = []
        buffer_len = 0

        for i, line in enumerate(lines):
            # Detect section headers
            header = self._detect_header(line)
            if header:
                # Flush current buffer before starting new section
                if buffer_lines:
                    flush_chunk(buffer_lines, current_section, paragraph_num)
                    # Overlap: carry last few lines into next chunk
                    overlap_text = "\n".join(buffer_lines)
                    overlap_chars = len(overlap_text)
                    if overlap_chars > self.valves.overlap:
                        overlap_lines = []
                        carried = 0
                        for bl in reversed(buffer_lines):
                            carried += len(bl) + 1
                            overlap_lines.insert(0, bl)
                            if carried >= self.valves.overlap:
                                break
                        buffer_lines = overlap_lines
                        buffer_len = sum(len(l) + 1 for l in buffer_lines)
                    else:
                        buffer_lines = []
                        buffer_len = 0
                    paragraph_num = 0

                current_section = header
                buffer_lines.append(line)
                buffer_len += len(line) + 1
                continue

            # Empty line = potential paragraph boundary
            if not line.strip():
                paragraph_num += 1
                buffer_lines.append(line)
                buffer_len += 1

                # Flush if buffer is large enough
                if buffer_len >= self.valves.chunk_size:
                    flush_chunk(buffer_lines, current_section, paragraph_num)
                    overlap_lines = []
                    carried = 0
                    for bl in reversed(buffer_lines):
                        carried += len(bl) + 1
                        overlap_lines.insert(0, bl)
                        if carried >= self.valves.overlap:
                            break
                    buffer_lines = overlap_lines
                    buffer_len = sum(len(l) + 1 for l in buffer_lines)
                continue

            buffer_lines.append(line)
            buffer_len += len(line) + 1

            # Hard split if a single paragraph exceeds chunk_size
            if buffer_len >= self.valves.chunk_size * 2:
                flush_chunk(buffer_lines, current_section, paragraph_num)
                paragraph_num += 1
                overlap_lines = []
                carried = 0
                for bl in reversed(buffer_lines):
                    carried += len(bl) + 1
                    overlap_lines.insert(0, bl)
                    if carried >= self.valves.overlap:
                        break
                buffer_lines = overlap_lines
                buffer_len = sum(len(l) + 1 for l in buffer_lines)

        # Final flush
        if buffer_lines:
            paragraph_num += 1
            flush_chunk(buffer_lines, current_section, paragraph_num)

        return chunks

    def _detect_header(self, line: str) -> Optional[str]:
        """
        Returns a normalized section label if the line looks like a header.
        Handles: '## Section Title', '3.', '3.1', 'ARTICLE IV', 'Section 4.', etc.
        """
        stripped = line.strip()
        if not stripped:
            return None

        # Markdown headers
        md = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if md:
            return md.group(2).strip()

        # Numbered section: 1., 1.1, 1.1.1, Section 3, Article IV
        num = re.match(
            r"^(?:Section|SECTION|Article|ARTICLE|Clause|CLAUSE)?\s*"
            r"(\d{1,3}(?:\.\d{1,3}){0,3}\.?|[IVXivx]{1,6}\.?)\s+"
            r"(.{3,60})$",
            stripped,
        )
        if num:
            title = num.group(2).strip()
            # Reject numbered list items (lowercase start or sentence-ending period)
            if title and title[0].isupper() and not title.endswith("."):
                return f"Section {num.group(1).rstrip('.')} — {title}"

        # All-caps short line (common in legal docs: "INDEMNIFICATION", "TERM AND TERMINATION")
        if (
            stripped.isupper()
            and 4 <= len(stripped) <= 60
            and not any(c.isdigit() for c in stripped[:3])
        ):
            return stripped.title()

        # Lines ending with colon that are short (e.g. "Payment Terms:")
        if stripped.endswith(":") and len(stripped) <= 60 and stripped[0].isupper():
            return stripped.rstrip(":")

        return None

    # ── Retrieval ──────────────────────────────────────────────────────────────

    def _retrieve(self, query: str, chunks: list, top_k: int) -> list:
        """
        Hybrid retrieval: BM25 keyword scores + semantic cosine similarity,
        fused with Reciprocal Rank Fusion.

        If no embedding model is available, falls back to pure BM25.
        Results are cached — chunk embeddings are computed once per document.
        """
        if self._embed_available:
            return self._retrieve_hybrid(query, chunks, top_k)
        else:
            return self._retrieve_bm25(query, chunks, top_k)

    def _retrieve_bm25(self, query: str, chunks: list, top_k: int) -> list:
        """
        BM25-style keyword retrieval. Scores each chunk by keyword overlap with
        the query, weighted by term frequency. No external dependencies.
        Returns top-k chunks sorted by score descending.
        """
        query_terms = self._tokenize(query)
        if not query_terms:
            return chunks[:top_k]

        # IDF-like weighting: terms appearing in fewer chunks are more discriminative
        doc_freq = {}
        for chunk in chunks:
            chunk_terms = set(self._tokenize(chunk["text"]))
            for term in query_terms:
                if term in chunk_terms:
                    doc_freq[term] = doc_freq.get(term, 0) + 1

        n_docs = max(len(chunks), 1)

        idf = {}
        for term in query_terms:
            df = doc_freq.get(term, 0)
            # BM25-style IDF: log((N - df + 0.5) / (df + 0.5) + 1)
            idf[term] = math.log((n_docs - df + 0.5) / (df + 0.5) + 1)

        scored = []
        for chunk in chunks:
            chunk_terms = self._tokenize(chunk["text"])
            term_counts = {}
            for t in chunk_terms:
                term_counts[t] = term_counts.get(t, 0) + 1

            chunk_len = len(chunk_terms)
            avg_len = 300  # approximate average chunk length in tokens
            k1, b = 1.5, 0.75

            score = 0.0
            for term in query_terms:
                tf = term_counts.get(term, 0)
                # BM25 TF normalization
                tf_norm = (tf * (k1 + 1)) / (
                    tf + k1 * (1 - b + b * chunk_len / avg_len)
                )
                score += idf.get(term, 0) * tf_norm

            # Boost if the section header contains query terms
            header_terms = set(self._tokenize(chunk.get("section", "")))
            header_hits = sum(1 for t in query_terms if t in header_terms)
            score += header_hits * 2.0

            if score > 0:
                scored.append({**chunk, "score": score})

        scored.sort(key=lambda x: x["score"], reverse=True)
        return scored[:top_k]

    def _retrieve_hybrid(self, query: str, chunks: list, top_k: int) -> list:
        """
        Hybrid retrieval using Reciprocal Rank Fusion over BM25 and semantic rankings.

        BM25 rank captures: exact legal terms, section numbers, dollar amounts.
        Semantic rank captures: conceptual queries ("bankruptcy" → "insolvency"),
          legal synonyms ("terminate" ↔ "cancel", "breach" ↔ "default"), and
          intent-level matches ("are we protected?" → indemnification clauses).

        Chunk embeddings are cached in self._embed_cache keyed by chunk text.
        The query is embedded fresh each call (queries are short — negligible cost).
        """
        # ── BM25 ranked list ──────────────────────────────────────────────────
        bm25_scored = self._retrieve_bm25(query, chunks, len(chunks))
        # For chunks with zero BM25 score, append them unscored so RRF can
        # still promote them via their semantic rank.
        bm25_seen = {c["_idx"] for c in bm25_scored}
        bm25_ranked = bm25_scored + [
            {**c, "score": 0.0} for c in chunks if c["_idx"] not in bm25_seen
        ]

        # ── Semantic ranked list ───────────────────────────────────────────────
        query_vec = _get_embedding(query, self._embed_model)
        if query_vec is None:
            # Embedding call failed at query time — degrade to BM25
            self._embed_available = False
            return bm25_scored[:top_k]

        # Embed each chunk (cached by text to avoid re-embedding on repeat queries)
        sem_scored = []
        for chunk in chunks:
            text = chunk["text"]
            if text not in self._embed_cache:
                vec = _get_embedding(text, self._embed_model)
                if vec is None:
                    # Partial failure — skip semantic for this chunk
                    continue
                self._embed_cache[text] = vec
            sim = _cosine_similarity(query_vec, self._embed_cache[text])
            sem_scored.append({**chunk, "score": sim})

        if not sem_scored:
            # All embedding calls failed
            self._embed_available = False
            return bm25_scored[:top_k]

        sem_scored.sort(key=lambda x: x["score"], reverse=True)

        # ── RRF fusion ─────────────────────────────────────────────────────────
        fused = _rrf_fuse(bm25_ranked, sem_scored, k=60)
        return fused[:top_k]

    # Legal term families: query with any member matches documents containing any member.
    # This handles morphological variants BM25 would otherwise miss.
    _LEGAL_SYNONYMS: dict = {
        "indemnif": ["indemnify", "indemnification", "indemnified", "indemnifies", "indemnitor", "indemnitee"],
        "terminat": ["terminate", "termination", "terminates", "terminated"],
        "arbitrat": ["arbitrate", "arbitration", "arbitrator", "arbitrated"],
        "confiden": ["confidential", "confidentiality", "confidentially"],
        "warrant": ["warranty", "warranties", "warrants", "warranted"],
        "represent": ["representation", "representations", "represent", "represents"],
        "liabil": ["liability", "liabilities", "liable"],
        "infringe": ["infringement", "infringe", "infringes", "infringed"],
        "assign": ["assignment", "assignee", "assignor", "assign", "assigns", "assigned"],
        "breach": ["breach", "breaches", "breached"],
    }
    # Map each surface form to its canonical root for normalization
    _LEGAL_STEM_MAP: dict = {}

    @classmethod
    def _build_stem_map(cls) -> None:
        if cls._LEGAL_STEM_MAP:
            return
        for root, forms in cls._LEGAL_SYNONYMS.items():
            for form in forms:
                cls._LEGAL_STEM_MAP[form] = root

    def _tokenize(self, text: str) -> list:
        """Lowercase, strip punctuation, split into tokens. Removes stopwords.
        Applies legal-domain stem normalization so morphological variants
        (e.g. 'indemnify' and 'indemnification') match each other."""
        stopwords = {
            "the", "a", "an", "and", "or", "but", "in", "on", "at", "to",
            "for", "of", "with", "by", "from", "is", "are", "was", "were",
            "be", "been", "being", "have", "has", "had", "do", "does", "did",
            "will", "would", "could", "should", "may", "might", "shall",
            "that", "this", "these", "those", "it", "its", "not", "no",
            "as", "up", "if", "out", "so", "than", "then", "such", "each",
            "any", "all", "both", "which", "who", "what", "when", "where",
        }
        self._build_stem_map()
        text = text.lower()
        tokens = re.findall(r"[a-z][a-z0-9'-]*[a-z0-9]|[a-z]", text)
        result = []
        for t in tokens:
            if t in stopwords or len(t) <= 1:
                continue
            # Normalize legal term families to their canonical root
            result.append(self._LEGAL_STEM_MAP.get(t, t))
        return result
